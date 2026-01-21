#!/usr/bin/env python3
"""
生成 DOCX 和 PDF 格式的测试合同文件

依赖安装：
pip install python-docx reportlab

运行本脚本：
python generate_contracts.py
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.enums import TA_LEFT
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

# 配置中文字体 - macOS 系统字体路径
CHINESE_FONT_PATHS = [
    '/System/Library/Fonts/STHeiti Light.ttc',  # 华文黑体
    '/System/Library/Fonts/PingFang.ttc',  # 苹方
    '/System/Library/Fonts/STSong.ttf',  # 华文宋体
    '/System/Library/Fonts/STKaiti.ttc',  # 华文楷体
]

# 全局字体名称
chinese_font_name = None

# 查找可用的中文字体
def register_chinese_font():
    """注册系统中可用的中文字体"""
    global chinese_font_name

    for font_path in CHINESE_FONT_PATHS:
        if os.path.exists(font_path):
            try:
                chinese_font_name = 'ChineseFont'
                pdfmetrics.registerFont(TTFont(chinese_font_name, font_path, subfontIndex=0))
                print(f"✓ 注册字体: {font_path}")
                return chinese_font_name
            except Exception as e:
                print(f"⚠️  字体注册失败 ({font_path}): {e}")
                continue

    print("⚠️  警告: 未找到可用的中文字体")
    return None

def create_docx_contract(txt_file, output_file):
    """从 TXT 文件创建 DOCX 文件"""
    doc = Document()

    # 读取文本内容
    with open(txt_file, 'r', encoding='utf-8') as f:
        content = f.read()

    # 添加标题
    filename = os.path.basename(txt_file).replace('.txt', '')
    doc.add_heading(filename, 0)

    # 添加段落
    paragraphs = content.split('\n\n')
    for para in paragraphs:
        lines = para.split('\n')
        for line in lines:
            if line.strip():
                doc.add_paragraph(line.strip())
        doc.add_paragraph()  # 段落间空行

    # 保存 DOCX
    doc.save(output_file)
    print(f"✓ 创建 DOCX: {output_file}")

def create_pdf_contract(txt_file, output_file):
    """从 TXT 文件创建 PDF 文件（使用 Platypus 以支持中文）"""
    # 读取文本内容
    with open(txt_file, 'r', encoding='utf-8') as f:
        content = f.read()

    # 创建 PDF 文档
    doc = SimpleDocTemplate(
        output_file,
        pagesize=A4,
        rightMargin=72,
        leftMargin=72,
        topMargin=72,
        bottomMargin=72
    )

    # 创建样式
    styles = getSampleStyleSheet()

    # 标题样式
    title_style = styles['Heading1']
    title_style.fontName = chinese_font_name or 'Helvetica'
    title_style.fontSize = 16
    title_style.leading = 20

    # 正文样式
    body_style = styles['BodyText']
    body_style.fontName = chinese_font_name or 'Helvetica'
    body_style.fontSize = 11
    body_style.leading = 16
    body_style.alignment = TA_LEFT

    # 构建内容
    story = []

    # 添加标题
    filename = os.path.basename(txt_file).replace('.txt', '')
    story.append(Paragraph(filename, title_style))
    story.append(Spacer(1, 12))

    # 添加段落内容
    paragraphs = content.split('\n\n')
    for para in paragraphs:
        if para.strip():
            # 将段落中的换行符替换为 <br/> 标签
            para_text = para.strip().replace('\n', '<br/>')
            story.append(Paragraph(para_text, body_style))
            story.append(Spacer(1, 6))

    # 生成 PDF
    try:
        doc.build(story)
        print(f"✓ 创建 PDF: {output_file}")
    except Exception as e:
        print(f"✗ 创建 PDF 失败 ({output_file}): {e}")

def main():
    """主函数"""
    # 定义合同文件映射
    contracts = [
        '技术服务合同.txt',
        '保密协议.txt',
        '劳动合同.txt',
        '销售合同.txt',
    ]

    print("开始生成测试合同文件...")
    print(f"当前目录: {os.getcwd()}")
    print()

    # 注册中文字体
    print("正在注册中文字体...")
    register_chinese_font()
    print()

    # 检查 TXT 文件是否存在
    for contract in contracts:
        if not os.path.exists(contract):
            print(f"⚠️  警告: 文件不存在 - {contract}")
            continue

        # 生成 DOCX
        docx_file = contract.replace('.txt', '.docx')
        try:
            create_docx_contract(contract, docx_file)
        except Exception as e:
            print(f"✗ 创建 DOCX 失败 ({contract}): {e}")

        # 生成 PDF
        pdf_file = contract.replace('.txt', '.pdf')
        try:
            create_pdf_contract(contract, pdf_file)
        except Exception as e:
            print(f"✗ 创建 PDF 失败 ({contract}): {e}")

    print()
    print("生成完成！")
    print()
    print("生成的文件：")
    for contract in contracts:
        if os.path.exists(contract.replace('.txt', '.docx')):
            print(f"  ✓ {contract.replace('.txt', '.docx')}")
        if os.path.exists(contract.replace('.txt', '.pdf')):
            print(f"  ✓ {contract.replace('.txt', '.pdf')}")

if __name__ == '__main__':
    main()
